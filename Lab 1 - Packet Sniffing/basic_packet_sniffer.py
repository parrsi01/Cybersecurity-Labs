# Name: Simon Parris
# Project: Packet Sniffer with Report
# Description:
#   Captures packets using Scapy, saves them into a timestamped .pcap file,
#   and generates a Word (.docx) summary report on exit.
#   Capture stops cleanly when you press Enter.

from scapy.all import sniff, wrpcap
from collections import Counter
from docx import Document
import datetime
import threading

# Global list to hold captured packets
captured_packets = []
stop_sniffing = threading.Event()

# -----------------------------
# Function: packet_callback
# -----------------------------
def packet_callback(packet):
    captured_packets.append(packet)

    print("\n--- Packet Captured ---")
    print(packet.summary())

    if packet.haslayer("IP"):
        ip_layer = packet["IP"]
        print(f"Source IP: {ip_layer.src}")
        print(f"Destination IP: {ip_layer.dst}")
        print(f"Protocol: {ip_layer.proto}")

    if packet.haslayer("TCP"):
        tcp_layer = packet["TCP"]
        print(f"TCP Source Port: {tcp_layer.sport}, TCP Destination Port: {tcp_layer.dport}")

    elif packet.haslayer("UDP"):
        udp_layer = packet["UDP"]
        print(f"UDP Source Port: {udp_layer.sport}, UDP Destination Port: {udp_layer.dport}")


# -----------------------------
# Function: generate_report
# Purpose: Create a Word document summarizing captured packets.
# -----------------------------
def generate_report(pcap_file, report_file):
    doc = Document()
    doc.add_heading("Packet Capture Analysis Report", 0)

    doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"PCAP File: {pcap_file}")
    doc.add_paragraph(f"Total Packets Captured: {len(captured_packets)}")

    # Count protocols
    protocols = Counter()
    src_ips = Counter()
    dst_ips = Counter()

    for pkt in captured_packets:
        if pkt.haslayer("IP"):
            ip_layer = pkt["IP"]
            src_ips[ip_layer.src] += 1
            dst_ips[ip_layer.dst] += 1
            protocols[ip_layer.proto] += 1

    # Protocol summary
    doc.add_heading("Protocol Summary", level=1)
    for proto, count in protocols.items():
        doc.add_paragraph(f"Protocol {proto}: {count} packets")

    # Top talkers
    doc.add_heading("Top Source IPs", level=1)
    for ip, count in src_ips.most_common(5):
        doc.add_paragraph(f"{ip}: {count} packets")

    doc.add_heading("Top Destination IPs", level=1)
    for ip, count in dst_ips.most_common(5):
        doc.add_paragraph(f"{ip}: {count} packets")

    doc.save(report_file)
    print(f"Report saved as {report_file}")


# -----------------------------
# Function: sniff_packets
# -----------------------------
def sniff_packets():
    def run_sniffer():
        sniff(iface=None, prn=packet_callback, store=False, stop_filter=lambda p: stop_sniffing.is_set())

    sniffer_thread = threading.Thread(target=run_sniffer)
    sniffer_thread.start()

    input("Press ENTER to stop the sniffer...\n")
    stop_sniffing.set()
    sniffer_thread.join()

    if captured_packets:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        pcap_file = f"captured_packets_{timestamp}.pcap"
        report_file = f"capture_report_{timestamp}.docx"

        print(f"\nSaving packets to {pcap_file}")
        wrpcap(pcap_file, captured_packets)

        print(f"Generating Word report: {report_file}")
        generate_report(pcap_file, report_file)
        print("Done!")
    else:
        print("No packets captured, nothing saved.")


# Entry point
if __name__ == "__main__":
    print("Starting packet sniffer... (Press ENTER to stop)")
    sniff_packets()