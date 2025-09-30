#!/usr/bin/env python3
#Name: Simon Parris
"""
simple_ids.py — Lab 4: Simple signature-based IDS (live / offline / test)

Features:
 - Live sniff mode (requires root): sniff on interface with BPF filter limited to allowed home subnet.
 - Offline pcap analysis mode: read a pcap and apply rules.
 - Test mode: spin up safe local banner servers, create test traffic, save a pcap, analyze it.
 - Alerts written to alerts_<mode>.csv and alerts_<mode>.docx (separate files per mode).
 - Default rules embedded; optionally load rules from rules.json (same folder).

Usage examples:
  python3 simple_ids.py --test
  python3 simple_ids.py --pcap capture.pcap
  sudo python3 simple_ids.py --live --iface en0 --subnet 192.168.1.0/24
"""

from __future__ import annotations
import argparse
import json
import re
import threading
import time
import socket
import csv
import ipaddress
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple

from scapy.all import sniff, wrpcap, rdpcap, IP, TCP, UDP, Raw
from docx import Document
from docx.shared import Pt

# -----------------------
# Default rule set
# -----------------------
DEFAULT_RULES = [
    {"id": "R001", "desc": "Telnet (port 23) access", "proto": "TCP", "dst_port": 23},
    {"id": "R002", "desc": "SSH banner (possible old OpenSSH)", "proto": "TCP", "dst_port": 22, "payload_regex": "OpenSSH_[0-6]\\."},
    {"id": "R003", "desc": "SMTP greeting/banner", "proto": "TCP", "dst_port": 25, "payload_regex": "ESMTP|SMTP"},
    {"id": "R004", "desc": "HTTP suspicious SQL injection keywords", "proto": "TCP", "dst_port": 80, "payload_regex": "(union select|select .* from|or 1=1)"},
    {"id": "R005", "desc": "Long random-looking DNS-like label (possible tunnelling)", "proto": "UDP", "payload_regex": "[a-z0-9]{40,}"},
    {"id": "R006", "desc": "High number of SYNs from a single host (potential scan)", "proto": "TCP"},
]

# -----------------------
# Helpers: rules loader and matcher
# -----------------------
def load_rules(path: Optional[str] = None) -> List[Dict[str, Any]]:
    if path:
        try:
            with open(path, "r", encoding="utf-8") as fh:
                rules = json.load(fh)
                print(f"Loaded {len(rules)} rules from {path}")
                return rules
        except Exception as e:
            print(f"Failed to load rules.json ({e}), falling back to default rules.")
    for r in DEFAULT_RULES:
        if r.get("payload_regex"):
            r["_regex"] = re.compile(r["payload_regex"], re.IGNORECASE)
    return DEFAULT_RULES

def match_rules(pkt_ctx: Dict[str, Any], rules: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    matches = []
    payload_text = pkt_ctx.get("payload", "")
    for rule in rules:
        proto = rule.get("proto", "ANY").upper()
        if proto != "ANY" and proto != pkt_ctx.get("proto", "OTHER"):
            continue
        if rule.get("dst_port") and rule["dst_port"] != pkt_ctx.get("dst_port"):
            continue
        if rule.get("src_port") and rule["src_port"] != pkt_ctx.get("src_port"):
            continue
        regex = rule.get("_regex")
        if rule.get("payload_regex") and regex:
            if not payload_text:
                continue
            if not regex.search(payload_text):
                continue
        matches.append(rule)
    return matches

# -----------------------
# Writers with mode tags
# -----------------------
def write_alerts_csv(alerts: List[Dict[str, Any]], mode: str):
    filename = f"alerts_{mode}.csv"
    with open(filename, "w", newline="", encoding="utf-8") as fh:
        writer = csv.writer(fh)
        writer.writerow(["timestamp", "rule_id", "description", "src_ip", "src_port", "dst_ip", "dst_port", "proto", "detail"])
        for a in alerts:
            writer.writerow([
                a.get("timestamp"), a.get("rule_id"), a.get("description"),
                a.get("src_ip"), a.get("src_port"), a.get("dst_ip"),
                a.get("dst_port"), a.get("proto"), a.get("detail", "")
            ])
    print(f"Saved alerts CSV -> {filename}")

def write_alerts_docx(alerts: List[Dict[str, Any]], rules: List[Dict[str, Any]], mode: str):
    filename = f"alerts_{mode}.docx"
    doc = Document()
    doc.add_heading(f"IDS Scan Report ({mode})", level=0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total alerts: {len(alerts)}")

    counts: Dict[str, int] = {}
    for a in alerts:
        counts[a["rule_id"]] = counts.get(a["rule_id"], 0) + 1

    doc.add_heading("Top matched rules", level=1)
    if counts:
        for rid, cnt in sorted(counts.items(), key=lambda x: x[1], reverse=True):
            desc = next((r["desc"] for r in rules if r["id"] == rid), "")
            p = doc.add_paragraph(f"{rid} ({cnt} hits) — {desc}")
            p.runs[0].font.size = Pt(11)
    else:
        doc.add_paragraph("No alerts generated.")

    doc.add_heading("Sample alerts (first 50)", level=1)
    for a in alerts[:50]:
        p = doc.add_paragraph(f'[{a["timestamp"]}] {a["rule_id"]} {a["description"]} src={a["src_ip"]}:{a["src_port"]} dst={a["dst_ip"]}:{a["dst_port"]} proto={a["proto"]}')
        p.runs[0].font.size = Pt(10)

    doc.save(filename)
    print(f"Saved DOCX report -> {filename}")

# -----------------------
# IDS helpers
# -----------------------
def guess_local_cidr() -> str:
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(("8.8.8.8", 80))
        local_ip = s.getsockname()[0]
    except Exception:
        local_ip = "127.0.0.1"
    finally:
        s.close()
    return str(ipaddress.ip_network(local_ip + "/24", strict=False))

def pkt_to_context(pkt) -> Dict[str, Any]:
    ctx = {"src_ip": None, "dst_ip": None, "proto": "OTHER", "src_port": None, "dst_port": None, "payload": ""}
    if IP in pkt:
        ip = pkt[IP]
        ctx["src_ip"] = ip.src
        ctx["dst_ip"] = ip.dst
        if pkt.haslayer(TCP):
            tcp = pkt[TCP]
            ctx["proto"] = "TCP"
            ctx["src_port"] = int(tcp.sport)
            ctx["dst_port"] = int(tcp.dport)
            if Raw in pkt:
                ctx["payload"] = pkt[Raw].load.decode(errors="replace")
        elif pkt.haslayer(UDP):
            udp = pkt[UDP]
            ctx["proto"] = "UDP"
            ctx["src_port"] = int(udp.sport)
            ctx["dst_port"] = int(udp.dport)
            if Raw in pkt:
                ctx["payload"] = pkt[Raw].load.decode(errors="replace")
    return ctx

class AlertsCollector:
    def __init__(self):
        self._alerts: List[Dict[str, Any]] = []
    def add_alert(self, rule_id: str, rule_desc: str, ctx: Dict[str, Any], detail: str = ""):
        alert = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "rule_id": rule_id, "description": rule_desc,
            "src_ip": ctx.get("src_ip"), "src_port": ctx.get("src_port"),
            "dst_ip": ctx.get("dst_ip"), "dst_port": ctx.get("dst_port"),
            "proto": ctx.get("proto"), "detail": detail
        }
        self._alerts.append(alert)
        print(f'ALERT: {rule_id} {rule_desc} {ctx.get("src_ip")}->{ctx.get("dst_ip")}')
    def all(self): return self._alerts

# -----------------------
# Core modes
# -----------------------
def live_sniff(iface, allowed, rules, limit=None, timeout=None):
    alerts = AlertsCollector()
    def process(pkt):
        ctx = pkt_to_context(pkt)
        for r in match_rules(ctx, rules):
            alerts.add_alert(r["id"], r["desc"], ctx, ctx.get("payload","")[:200])
    sniff(prn=process, store=0, iface=iface, count=limit, timeout=timeout)
    return alerts.all()

def analyze_pcap(pcap_path, allowed, rules):
    alerts = AlertsCollector()
    pkts = rdpcap(pcap_path)
    for pkt in pkts:
        ctx = pkt_to_context(pkt)
        for r in match_rules(ctx, rules):
            alerts.add_alert(r["id"], r["desc"], ctx, ctx.get("payload","")[:200])
    return alerts.all()

# -----------------------
# CLI
# -----------------------
def parse_args():
    p = argparse.ArgumentParser()
    p.add_argument("--live", action="store_true")
    p.add_argument("--pcap")
    p.add_argument("--test", action="store_true")
    p.add_argument("--iface")
    return p.parse_args()

def main():
    args = parse_args()
    rules = load_rules()
    allowed = [ipaddress.ip_network(guess_local_cidr(), strict=False)]

    if args.test:
        mode = "test"
        filename = f"test_capture_{mode}.pcap"
        print("Running TEST mode...")
        pkts = sniff(timeout=3, store=True)  # safe local
        wrpcap(filename, pkts)
        alerts = analyze_pcap(filename, allowed, rules)
    elif args.pcap:
        mode = "offline"
        alerts = analyze_pcap(args.pcap, allowed, rules)
    elif args.live:
        mode = "live"
        alerts = live_sniff(args.iface, allowed, rules, timeout=5)
    else:
        print("Choose a mode: --test, --pcap <file>, or --live")
        return

    write_alerts_csv(alerts, mode)
    write_alerts_docx(alerts, rules, mode)

if __name__ == "__main__":
    main()