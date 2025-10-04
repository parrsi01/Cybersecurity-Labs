#!/usr/bin/env python3
# Name: Simon Parris
"""
Lab 5 — Anomaly-based IDS
Detects unusual packets by size, port, or protocol.
Automatically saves results to CSV, DOCX, and PCAP.
"""

import argparse
import csv
import os
from datetime import datetime
from scapy.all import sniff, rdpcap, wrpcap, IP, TCP, UDP
from docx import Document

# -----------------------
# Detection rules
# -----------------------
def detect_anomaly(pkt):
    """Check for unusual patterns in packet headers."""
    if not pkt.haslayer(IP):
        return (False, "non-IP")

    plen = len(pkt)
    proto = pkt[IP].proto
    sport, dport = None, None
    if pkt.haslayer(TCP):
        sport, dport = pkt[TCP].sport, pkt[TCP].dport
    elif pkt.haslayer(UDP):
        sport, dport = pkt[UDP].sport, pkt[UDP].dport

    # Simple anomaly thresholds
    if plen < 40:
        return (True, f"Suspiciously small packet ({plen} bytes)")
    if plen > 1200:
        return (True, f"Suspiciously large packet ({plen} bytes)")
    if sport and sport > 50000:
        return (True, f"Unusual source port {sport}")
    if dport and dport > 50000:
        return (True, f"Unusual destination port {dport}")
    if proto not in [6, 17]:  # 6 = TCP, 17 = UDP
        return (True, f"Unusual protocol {proto}")

    return (False, "normal")


# -----------------------
# Analyze packets
# -----------------------
def analyze_packets(pkts):
    alerts, anomalous_pkts = [], []
    for i, pkt in enumerate(pkts, 1):
        is_anom, reason = detect_anomaly(pkt)
        if is_anom:
            alerts.append({
                "index": i,
                "src": pkt[IP].src if pkt.haslayer(IP) else "?",
                "dst": pkt[IP].dst if pkt.haslayer(IP) else "?",
                "reason": reason
            })
            anomalous_pkts.append(pkt)
    return alerts, anomalous_pkts


# -----------------------
# Save outputs
# -----------------------
def save_csv(alerts, filename):
    with open(filename, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["index", "src", "dst", "reason"])
        writer.writeheader()
        writer.writerows(alerts)
    print(f"[+] CSV saved -> {filename}")


def save_docx(alerts, filename):
    doc = Document()
    doc.add_heading("Anomaly-based IDS Report", 0)
    doc.add_paragraph(f"Total anomalies detected: {len(alerts)}")
    for a in alerts:
        doc.add_paragraph(
            f"Packet {a['index']} | {a['src']} → {a['dst']} | Reason: {a['reason']}"
        )
    doc.save(filename)
    print(f"[+] DOCX saved -> {filename}")


def save_pcap(pkts, filename):
    if pkts:
        wrpcap(filename, pkts)
        print(f"[+] PCAP saved -> {filename}")
    else:
        print("[i] No anomalies detected — no PCAP written.")


# -----------------------
# Main logic
# -----------------------
def main():
    parser = argparse.ArgumentParser(description="Anomaly-based Intrusion Detection System")
    parser.add_argument("--live", action="store_true", help="Capture live packets")
    parser.add_argument("--iface", help="Interface for live mode (e.g., en0, eth0)")
    parser.add_argument("--pcap", help="Input PCAP file to analyze")
    parser.add_argument("--count", type=int, default=50, help="Packets to capture (live mode only)")
    args = parser.parse_args()

    # File naming scheme (unique per run)
    base_name = "anomaly_results"
    csv_file = f"{base_name}.csv"
    docx_file = f"{base_name}.docx"
    pcap_file = f"{base_name}.pcap"

    if args.live:
        if not args.iface:
            print("Error: must specify --iface for live mode.")
            return
        print(f"[*] Capturing {args.count} packets on {args.iface} ...")
        pkts = sniff(iface=args.iface, count=args.count)
    elif args.pcap:
        print(f"[*] Reading packets from {args.pcap}")
        pkts = rdpcap(args.pcap)
    else:
        print("Error: specify --live or --pcap")
        return

    print(f"[*] Analyzing {len(pkts)} packets ...")
    alerts, anomalies = analyze_packets(pkts)

    save_csv(alerts, csv_file)
    save_docx(alerts, docx_file)
    save_pcap(anomalies, pcap_file)
    print("[✓] Analysis complete.")


if __name__ == "__main__":
    main()