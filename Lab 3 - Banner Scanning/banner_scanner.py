#!/usr/bin/env python3
#Name: Simon Parris
"""
banner_scanner.py

Banner grabbing + safe local test mode with PCAP capture and DOCX report.

Features:
 - Captures packets (PCAP) for test/scan runs (no timestamps in filenames).
 - Writes DOCX summary report for each run mode (no timestamps).
 - CSV export option remains available.
 - Uses AsyncSniffer from scapy to capture traffic during scans.

IMPORTANT: packet capture often requires elevated privileges (sudo).
Only scan hosts/networks you own or have permission to test.
"""

from __future__ import annotations
import socket
import argparse
import threading
import csv
import os
from http.server import HTTPServer, SimpleHTTPRequestHandler
from typing import List, Tuple, Optional

# Scapy for packet capture
from scapy.all import AsyncSniffer, wrpcap

# docx reporting
try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None

# -----------------------
# Defaults
# -----------------------
DEFAULT_PORTS = [21, 22, 25, 80, 443, 3306, 6379]
DEFAULT_TIMEOUT = 2.0

PORT_PROBES = {
    21: None,
    22: None,
    25: b"HELO example.com\r\n",
    80: b"GET / HTTP/1.0\r\nHost: example.com\r\n\r\n",
    443: b"HEAD / HTTP/1.0\r\nHost: example.com\r\n\r\n",
    3306: None,
    6379: b"PING\r\n",
}

# -----------------------
# Filenames (no timestamps)
# -----------------------
PCAP_TEST = "banners_test.pcap"
PCAP_OFFLINE = "banners_offline.pcap"
PCAP_LIVE = "banners_live.pcap"
DOCX_TEST = "banners_test.docx"
DOCX_OFFLINE = "banners_offline.docx"
DOCX_LIVE = "banners_live.docx"

# -----------------------
# Banner grabber
# -----------------------
def grab_banner(host: str, port: int, probe: Optional[bytes], timeout: float) -> str:
    """
    Connect to (host, port), optionally send probe, and attempt to read reply.
    Returns a human-friendly string describing the result.
    """
    try:
        with socket.create_connection((host, port), timeout=timeout) as s:
            s.settimeout(timeout)
            if probe:
                try:
                    s.sendall(probe)
                except Exception:
                    pass

            try:
                data = s.recv(4096)
                if not data:
                    return "<connected but no banner data received>"
                return data.decode(errors="replace").strip().replace("\r\n", " | ")
            except socket.timeout:
                return "<connected but read timed out>"
            except Exception as e:
                return f"<error reading: {e}>"
    except Exception as e:
        return f"<connect error: {e}>"

# -----------------------
# CSV helper
# -----------------------
def save_csv(filename: str, host: str, results: List[Tuple[int, str]]):
    with open(filename, "w", newline="", encoding="utf-8") as fh:
        writer = csv.writer(fh)
        writer.writerow(["host", "port", "banner"])
        for port, banner in results:
            writer.writerow([host, port, banner])
    print(f"CSV written -> {filename}")

# -----------------------
# DOCX helper
# -----------------------
def write_docx(results: List[Tuple[int, str]], filename: str, title: str = "Banner Scan Report"):
    if Document is None:
        print("[!] python-docx not installed; skipping DOCX output.")
        return
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Total results: {len(results)}")
    doc.add_heading("Findings", level=1)
    for port, banner in results:
        p = doc.add_paragraph()
        p.add_run(f"Port {port}:").bold = True
        p.add_run(f" {banner}\n")
        p.runs[0].font.size = Pt(11)
    doc.save(filename)
    print(f"DOCX saved -> {filename}")

# -----------------------
# Safe test servers (local)
# -----------------------
def start_http_server(port: int):
    handler = SimpleHTTPRequestHandler
    httpd = HTTPServer(("127.0.0.1", port), handler)
    thread = threading.Thread(target=httpd.serve_forever, daemon=True)
    thread.start()
    return httpd, thread

def _banner_tcp_server(port: int, banner_bytes: bytes, stop_event: threading.Event):
    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
    sock.bind(("127.0.0.1", port))
    sock.listen(5)
    sock.settimeout(1.0)
    try:
        while not stop_event.is_set():
            try:
                client, addr = sock.accept()
                try:
                    client.sendall(banner_bytes)
                except Exception:
                    pass
                try:
                    client.shutdown(socket.SHUT_WR)
                except Exception:
                    pass
                client.close()
            except socket.timeout:
                continue
    finally:
        sock.close()

def start_banner_server(port: int, banner: str) -> threading.Event:
    event = threading.Event()
    t = threading.Thread(target=_banner_tcp_server, args=(port, banner.encode(), event), daemon=True)
    t.start()
    return event

# -----------------------
# Capture wrapper using AsyncSniffer
# -----------------------
def start_capture(host: str, iface: Optional[str] = None) -> AsyncSniffer:
    """
    Start an AsyncSniffer capturing only traffic to/from host.
    Returns the sniffer object (must call stop() and collect results).
    """
    bpf = f"host {host}"
    sniffer = AsyncSniffer(filter=bpf, iface=iface, store=True)
    sniffer.start()
    return sniffer

# -----------------------
# Orchestration / scanning
# -----------------------
def run_scan(host: str, ports: List[int], timeout: float, save_csv_opt: Optional[str], capture_pcap: Optional[str], iface: Optional[str]) -> List[Tuple[int, str]]:
    results: List[Tuple[int, str]] = []
    sniffer = None
    try:
        if capture_pcap:
            try:
                sniffer = start_capture(host, iface=iface)
                print(f"Packet capture started for host {host} (filter: host {host})")
            except Exception as e:
                print(f"[!] Could not start packet capture: {e}")

        for port in ports:
            probe = PORT_PROBES.get(port, None)
            banner = grab_banner(host, port, probe, timeout)
            print(f"{host}:{port} -> {banner}")
            results.append((port, banner))

        # stop capture and save pcap
        if sniffer:
            try:
                captured = sniffer.stop()
                wrpcap(capture_pcap, captured)
                print(f"Saved PCAP -> {capture_pcap} ({len(captured)} packets)")
            except Exception as e:
                print(f"[!] Error saving pcap: {e}")

    finally:
        if save_csv_opt:
            save_csv(save_csv_opt, host, results)
    return results

# -----------------------
# CLI / main
# -----------------------
def parse_ports_arg(s: Optional[str]) -> List[int]:
    if not s:
        return DEFAULT_PORTS.copy()
    parts = s.split(",")
    out: List[int] = []
    for p in parts:
        p = p.strip()
        if "-" in p:
            a, b = p.split("-", 1)
            try:
                a_i = int(a); b_i = int(b)
                out.extend(list(range(a_i, b_i + 1)))
            except ValueError:
                continue
        else:
            try:
                out.append(int(p))
            except ValueError:
                continue
    return sorted(set(out))

def main():
    parser = argparse.ArgumentParser(description="Banner scanner with PCAP capture and DOCX outputs.")
    parser.add_argument("host", nargs="?", default="localhost", help="Target host (default localhost).")
    parser.add_argument("--ports", help="Comma list or ranges, e.g. 22,80,8000-8010. Defaults to common ports.", default=None)
    parser.add_argument("--timeout", type=float, default=DEFAULT_TIMEOUT, help="Socket timeout seconds.")
    parser.add_argument("--save-csv", metavar="FILE", help="Save results to CSV file (optional).")
    parser.add_argument("--test", action="store_true", help="Start safe local test servers and scan localhost.")
    parser.add_argument("--capture-iface", help="Optional interface to capture on (e.g., en0). If omitted, default is used.")
    args = parser.parse_args()

    if args.test:
        print("** TEST MODE ** Starting local test servers (HTTP + 2 banner TCP servers)...")
        http_port = 8080
        httpd, _ = start_http_server(http_port)
        ssh_port = 2222
        smtp_port = 2525
        ssh_event = start_banner_server(ssh_port, "SSH-2.0-OpenSSH_TestServer_1.0\r\n")
        smtp_event = start_banner_server(smtp_port, "220 localhost ESMTP TestMailService\r\n")
        host = "127.0.0.1"
        ports = [ssh_port, smtp_port, http_port]
        pcap_name = PCAP_TEST
        docx_name = DOCX_TEST
    else:
        host = args.host
        ports = parse_ports_arg(args.ports)
        pcap_name = PCAP_OFFLINE
        docx_name = DOCX_OFFLINE

    print(f"Target: {host}")
    print(f"Ports: {ports}")

    try:
        capture_iface = args.capture_iface
        capture_pcap = pcap_name
        results = run_scan(host, ports, args.timeout, args.save_csv, capture_pcap, capture_iface)

        # Write DOCX report (no timestamp in filename)
        write_docx(results, docx_name, title=f"Banner Scan Report ({'test' if args.test else 'scan'})")

    finally:
        if args.test:
            print("Shutting down test servers...")
            try:
                ssh_event.set()
                smtp_event.set()
            except Exception:
                pass
            try:
                httpd.shutdown()
            except Exception:
                pass

    print("Done.")

if __name__ == "__main__":
    main()