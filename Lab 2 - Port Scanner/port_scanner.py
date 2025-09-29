"""
port_scanner.py
Author: Simon Parris
Project: Lab 2 - Port Scanner (Cybersecurity Labs)

Features:
 - Scans a TCP port range and prints results to console.
 - Optional CSV detailed output (--save-csv) -> writes to portscan.csv
 - Optional Word (.docx) report (--save-docx) -> writes to portscan.docx
 - NOTE: PCAP output is NOT applicable for a port scanner.
"""

import socket
import argparse
from datetime import datetime
from typing import List, Tuple
import csv
from docx import Document
from docx.shared import Pt

# -----------------------
# Configuration defaults
# -----------------------
DEFAULT_START_PORT = 20
DEFAULT_END_PORT = 1024
SOCKET_TIMEOUT = 0.5


# -----------------------
# Scanning primitives
# -----------------------
def scan_port(host: str, port: int) -> bool:
    """
    Attempt a TCP connect to (host, port). Returns True if open.
    """
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        s.settimeout(SOCKET_TIMEOUT)
        result = s.connect_ex((host, port))
        s.close()
        return result == 0
    except Exception:
        return False


def scan_range(host: str, start_port: int, end_port: int) -> List[Tuple[int, bool]]:
    """
    Scan ports in range and return list of (port, is_open).
    Prints per-port status while scanning.
    """
    print(f"\n🔎 Scanning {host} ports {start_port}..{end_port}\n")
    results: List[Tuple[int, bool]] = []

    for port in range(start_port, end_port + 1):
        is_open = scan_port(host, port)
        results.append((port, is_open))
        if is_open:
            print(f"✅ Port {port:5d} OPEN")
        else:
            print(f"❌ Port {port:5d} closed")

    print("\n📊 Scan finished.")
    return results


# -----------------------
# Output helpers (no timestamps in filenames)
# -----------------------
def save_csv(host: str, start_port: int, end_port: int, results: List[Tuple[int, bool]]) -> str:
    """
    Save per-port results to CSV. Returns filename.
    Overwrites portscan.csv if it exists.
    """
    filename = "portscan.csv"
    with open(filename, "w", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(["target", "start_port", "end_port", "generated_at"])
        writer.writerow([host, start_port, end_port, datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
        writer.writerow([])  # blank line
        writer.writerow(["port", "is_open"])
        for port, is_open in results:
            writer.writerow([port, int(is_open)])  # 1=open, 0=closed
    print(f"CSV saved -> {filename}")
    return filename


def save_docx(host: str, start_port: int, end_port: int, results: List[Tuple[int, bool]]) -> str:
    """
    Generate a Word (.docx) report summarizing the scan and listing open ports.
    Overwrites portscan.docx if it exists.
    """
    filename = "portscan.docx"
    doc = Document()
    title = f"Port Scan Report — {host}"
    doc.add_heading(title, level=0)

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Target: {host}")
    doc.add_paragraph(f"Port range: {start_port} - {end_port}")

    total = len(results)
    open_ports = [port for port, open_flag in results if open_flag]
    doc.add_paragraph(f"Total ports scanned: {total}")
    doc.add_paragraph(f"Open ports count: {len(open_ports)}")

    doc.add_heading("Open ports", level=1)
    if open_ports:
        for p in open_ports:
            p_par = doc.add_paragraph(str(p))
            p_par.runs[0].font.size = Pt(11)
    else:
        doc.add_paragraph("None found")

    # Add a brief table of first N results to show example rows
    doc.add_heading("Sample results (first 30 ports scanned)", level=1)
    sample = results[:30]
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Port"
    hdr_cells[1].text = "Open (1=yes,0=no)"
    for port, is_open in sample:
        row_cells = table.add_row().cells
        row_cells[0].text = str(port)
        row_cells[1].text = "1" if is_open else "0"

    doc.save(filename)
    print(f"Docx report saved -> {filename}")
    return filename


# -----------------------
# CLI parsing & main
# -----------------------
def parse_args():
    parser = argparse.ArgumentParser(description="TCP port scanner with optional CSV/DOCX export")
    parser.add_argument("host", help="Target hostname or IPv4 address")
    parser.add_argument("--start", type=int, default=DEFAULT_START_PORT, help=f"Start port (default {DEFAULT_START_PORT})")
    parser.add_argument("--end", type=int, default=DEFAULT_END_PORT, help=f"End port (default {DEFAULT_END_PORT})")
    parser.add_argument("--save-csv", action="store_true", help="Save detailed CSV of results (portscan.csv)")
    parser.add_argument("--save-docx", action="store_true", help="Save Word (.docx) summary report (portscan.docx)")
    return parser.parse_args()


def main():
    args = parse_args()
    start_port = max(1, min(65535, args.start))
    end_port = max(1, min(65535, args.end))
    if start_port > end_port:
        start_port, end_port = end_port, start_port

    results = scan_range(args.host, start_port, end_port)

    # Summary printed to console
    open_ports = [p for p, flag in results if flag]
    print("\n--- Scan Summary ---")
    print(f"Target: {args.host}")
    print(f"Range scanned: {start_port} - {end_port} ({len(results)} ports)")
    print(f"Open ports: {', '.join(map(str, open_ports)) if open_ports else 'None found'}")

    # Optional outputs (no timestamps in filenames)
    if args.save_csv:
        save_csv(args.host, start_port, end_port, results)
    if args.save_docx:
        save_docx(args.host, start_port, end_port, results)

    if not args.save_csv and not args.save_docx:
        print("\nTip: Use --save-csv or --save-docx to save results. PCAP is not applicable to this lab.")


if __name__ == "__main__":
    main()